#!/usr/bin/env python3
"""
Gera uma petição .docx EDITÁVEL a partir de um Markdown, replicando o padrão
das peças do escritório (calibrado sobre modelo real ROSELI DA ROSA SCRINZ).

Uso:
    python3 scripts/gerar_docx.py entrada.md saida.docx

Padrão aplicado:
- A4, margens 2,0 cm em todos os lados
- Times New Roman 12, JUSTIFICADO, entrelinha 1,5, sem espaço antes/depois
- Recuo de 1ª linha de 1,5 cm no corpo E nos títulos de tópico (títulos em negrito)
- Endereçamento (EXCELENTÍSSIMO/AO JUÍZO...): negrito, justificado, sem recuo
- "RECLAMAÇÃO TRABALHISTA" ou "CONTESTAÇÃO" (linha isolada): centralizado, negrito
- Cabeçalho dos autos ("Processo nº:", "Rito:", "Reclamante:", "Reclamada:"):
  justificado, sem recuo
- Ementas/citações (linhas iniciadas por ">"): recuo à esquerda de 4 cm
- Cidade/data, nome do advogado e OAB: centralizados em negrito
- Separação entre blocos por linha em branco (preservada como parágrafo vazio)
- **negrito** inline; tabelas Markdown viram tabelas do Word

Instala o python-docx automaticamente se faltar.
"""
import sys, os, re, subprocess

try:
    import docx
except ImportError:
    subprocess.run([sys.executable, "-m", "pip", "install", "--quiet", "python-docx"], check=True)
    import docx

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

FONTE, CORPO = "Times New Roman", 12
RECUO = Cm(1.5)
RECUO_EMENTA = Cm(4.0)

# Papel timbrado do escritório (logo no cabeçalho + barra de contato no rodapé).
# Procurado em assets/ na raiz do projeto (um nível acima de scripts/).
_ASSETS = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), "assets")
LOGO = os.path.join(_ASSETS, "logo-cabecalho.png")   # 5,97 x 3,49 cm
RODAPE = os.path.join(_ASSETS, "rodape.png")          # 17,0 x 2,64 cm


def add_papel_timbrado(document):
    """Adiciona logo (cabeçalho) e barra de contato (rodapé), se os arquivos existirem."""
    sec = document.sections[0]
    sec.header_distance = Cm(1.0)
    sec.footer_distance = Cm(1.0)
    if os.path.exists(LOGO):
        hp = sec.header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        hp.paragraph_format.first_line_indent = Cm(0)
        hp.add_run().add_picture(LOGO, width=Cm(5.97))
        sec.header.is_linked_to_previous = False
    if os.path.exists(RODAPE):
        fp = sec.footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fp.paragraph_format.first_line_indent = Cm(0)
        fp.add_run().add_picture(RODAPE, width=Cm(17.0))
        sec.footer.is_linked_to_previous = False

RE_ENDERECO = re.compile(r"^\s*(EXCELENT[IÍ]SSIM|EXM|AO JU[IÍ]ZO|MERIT[IÍ]SSIM)", re.IGNORECASE)
RE_TITULO = re.compile(r"^\s*(RECLAMA(ÇÃO|CAO)\s+TRABALHISTA|CONTESTA(ÇÃO|CAO))\s*$", re.IGNORECASE)
# Cabeçalho dos autos na contestação (Processo nº, Rito, Reclamante, Reclamada):
# justificado, sem recuo de 1ª linha, sem negrito.
RE_AUTOS = re.compile(r"^\s*(Processo|Autos|Rito|Reclamante|Reclamad[oa]|Reclamad[oa]\s*\d|R[eé]u|R[eé])\s*(n?º|nº|n\.º)?\s*:", re.IGNORECASE)
RE_CIDADE_DATA = re.compile(r"^\s*[\wÀ-ú\.\s]{2,40}/[A-Z]{2},\s*(data do protocolo|\d{1,2}\s+de\s+\w+\s+de\s+\d{4}|data\b).*$", re.IGNORECASE)
RE_OAB = re.compile(r"OAB", re.IGNORECASE)
BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def _font(run, size=CORPO, bold=False, italic=False):
    run.font.name = FONTE
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), FONTE)


def set_base_style(document):
    st = document.styles["Normal"]
    st.font.name = FONTE
    st.font.size = Pt(CORPO)
    rf = st.element.get_or_add_rPr().get_or_add_rFonts()
    for a in ("w:ascii", "w:hAnsi", "w:cs"):
        rf.set(qn(a), FONTE)
    pf = st.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    for s in document.sections:
        s.page_height, s.page_width = Cm(29.7), Cm(21.0)
        s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.0)


VERMELHO = RGBColor(0xFF, 0x00, 0x00)
# tokens inline: **negrito** ou [placeholder a preencher/conferir]
TOKEN = re.compile(r"\*\*(.+?)\*\*|\[[^\]]*\]")


def add_runs(paragraph, text, size=CORPO, base_bold=False, italic=False):
    """Adiciona texto interpretando **negrito** e deixando [colchetes] em VERMELHO."""
    pos = 0
    added = False
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            _font(paragraph.add_run(text[pos:m.start()]), size=size, bold=base_bold, italic=italic)
        tok = m.group(0)
        if tok.startswith("**"):
            inner = m.group(1)
            r = paragraph.add_run(inner)
            _font(r, size=size, bold=True, italic=italic)
            if inner.strip().startswith("[") and inner.strip().endswith("]"):
                r.font.color.rgb = VERMELHO  # placeholder em negrito -> negrito + vermelho
        else:  # [ ... ] -> campo a preencher / conferir -> vermelho
            r = paragraph.add_run(tok)
            _font(r, size=size, bold=base_bold, italic=italic)
            r.font.color.rgb = VERMELHO
        pos = m.end()
        added = True
    if pos < len(text) or not added:
        _font(paragraph.add_run(text[pos:]), size=size, bold=base_bold, italic=italic)


def is_heading(s):
    if s.startswith("#"):
        return True
    letters = [c for c in s if c.isalpha()]
    return bool(letters) and len(s) <= 130 and len(letters) >= 3 and \
        sum(1 for c in letters if c.isupper()) / len(letters) >= 0.9


def parse_table(lines, i):
    rows = []
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(set(c) <= set("-: ") for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def next_nonempty(lines, i):
    j = i + 1
    while j < len(lines) and not lines[j].strip():
        j += 1
    return lines[j].strip() if j < len(lines) else ""


def main():
    if len(sys.argv) < 3:
        print("Uso: python3 gerar_docx.py entrada.md saida.docx"); sys.exit(1)
    entrada, saida = sys.argv[1], sys.argv[2]
    with open(entrada, encoding="utf-8") as f:
        lines = f.read().split("\n")

    doc = Document()
    set_base_style(doc)
    add_papel_timbrado(doc)
    pending_blank = False
    wrote = False

    i = 0
    while i < len(lines):
        raw = lines[i]
        s = raw.strip()

        if not s:
            pending_blank = True
            i += 1
            continue

        if s.startswith("|"):
            if pending_blank and wrote:
                doc.add_paragraph()
            pending_blank = False
            rows, i = parse_table(lines, i)
            if rows:
                ncols = max(len(r) for r in rows)
                t = doc.add_table(rows=0, cols=ncols)
                t.style = "Table Grid"
                t.alignment = WD_TABLE_ALIGNMENT.CENTER
                for r, cells in enumerate(rows):
                    cells += [""] * (ncols - len(cells))
                    rc = t.add_row().cells
                    for c, val in enumerate(cells):
                        para = rc[c].paragraphs[0]
                        para.paragraph_format.first_line_indent = Cm(0)
                        para.paragraph_format.line_spacing = 1.0
                        add_runs(para, val, size=11, base_bold=(r == 0))
                wrote = True
            continue

        if pending_blank and wrote:
            doc.add_paragraph()
        pending_blank = False

        p = doc.add_paragraph()
        pf = p.paragraph_format
        text = re.sub(r"^#+\s*", "", s)
        J, C = WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.CENTER

        if RE_ENDERECO.match(s):
            p.alignment = J; pf.first_line_indent = Cm(0)
            add_runs(p, text, base_bold=True)
        elif RE_AUTOS.match(s):
            p.alignment = J; pf.first_line_indent = Cm(0)
            add_runs(p, text)
        elif RE_TITULO.match(text):
            p.alignment = C; pf.first_line_indent = Cm(0)
            add_runs(p, text, base_bold=True)
        elif RE_CIDADE_DATA.match(s) or RE_OAB.search(s) or (is_heading(s) and RE_OAB.search(next_nonempty(lines, i))):
            # bloco de assinatura: cidade/data, nome do advogado, OAB
            p.alignment = C; pf.first_line_indent = Cm(0)
            add_runs(p, text, base_bold=True)
        elif s.startswith(">"):
            p.alignment = J; pf.first_line_indent = Cm(0); pf.left_indent = RECUO_EMENTA
            add_runs(p, re.sub(r"^>\s?", "", s))
        elif is_heading(s):
            p.alignment = J; pf.first_line_indent = RECUO
            add_runs(p, text, base_bold=True)
        else:
            p.alignment = J; pf.first_line_indent = RECUO
            add_runs(p, text)
        wrote = True
        i += 1

    doc.save(saida)
    print(f"OK: {saida}")


if __name__ == "__main__":
    main()
